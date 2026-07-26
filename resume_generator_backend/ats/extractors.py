"""Stage 1 — text extraction. PDFs get dual extraction (pdfplumber and
PyMuPDF) so downstream checks can compare the two; DOCX has a single
extractor and the agreement check auto-passes."""

import io
import re

import docx
import fitz  # PyMuPDF
import pdfplumber

from ats.thresholds import EDGE_PROXIMITY_FRACTION, TINY_FONT_PT

# Embedded font subsets are named "ABCDEF+RealName"; strip the tag so the
# same font subsetted twice is not counted as two fonts.
_SUBSET_PREFIX_RE = re.compile(r"^[A-Z]{6}\+")
# LMRoman10-Regular-Identity-H and LMRoman12-Regular-Identity-H are the SAME
# family at different optical sizes (standard pdflatex output) — normalize to
# the family so font-count measures families, not size/style variants.
_FONT_NOISE_RE = re.compile(
    r"(-Identity-H|-(Regular|Bold|Italic|BoldItalic|Oblique|Light|Medium))+$",
    re.IGNORECASE,
)
_FONT_SIZE_SUFFIX_RE = re.compile(r"\d+$")


def _font_family(basefont: str) -> str:
    name = _SUBSET_PREFIX_RE.sub("", basefont)
    name = _FONT_NOISE_RE.sub("", name)
    return _FONT_SIZE_SUFFIX_RE.sub("", name) or name


def extract_pdf_pdfplumber(data):
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        return "\n".join((page.extract_text() or "") for page in pdf.pages)


def extract_pdf_pymupdf(data):
    with fitz.open(stream=data, filetype="pdf") as doc:
        return "\n".join(page.get_text() for page in doc)


def extract_pdf_links(data):
    """URI targets of every link annotation, across all pages."""
    with fitz.open(stream=data, filetype="pdf") as doc:
        return [
            link["uri"] for page in doc for link in page.get_links() if link.get("uri")
        ]


def pdf_stats(data):
    """Document-shape stats for the extended check battery.

    One pass over the PyMuPDF document (pages, fonts, glyph sizes, images)
    plus one over pdfplumber (word bboxes near the page edges). Returns
    {page_count, font_names, tiny_char_fraction, image_count, edge_text,
    is_encrypted}.
    """
    font_names = set()
    image_xrefs = set()
    tiny_chars = 0
    total_chars = 0
    with fitz.open(stream=data, filetype="pdf") as doc:
        page_count = doc.page_count
        # Owner-password ("permissions") PDFs auto-authenticate with the
        # empty user password, leaving is_encrypted/needs_pass False — the
        # metadata encryption method string still exposes them.
        is_encrypted = bool(
            doc.is_encrypted or doc.needs_pass or (doc.metadata or {}).get("encryption")
        )
        for page in doc:
            for font in page.get_fonts(full=True):
                basefont = font[3]
                if basefont:
                    font_names.add(_font_family(basefont))
            for image in page.get_images(full=True):
                image_xrefs.add(image[0])  # xref — same image reused counts once
            for block in page.get_text("dict")["blocks"]:
                for line in block.get("lines", ()):
                    for span in line["spans"]:
                        count = len(span["text"])
                        total_chars += count
                        if span["size"] < TINY_FONT_PT:
                            tiny_chars += count

    edge_text = False
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            margin_x = page.width * EDGE_PROXIMITY_FRACTION
            margin_y = page.height * EDGE_PROXIMITY_FRACTION
            if any(
                word["x0"] < margin_x
                or word["x1"] > page.width - margin_x
                or word["top"] < margin_y
                or word["bottom"] > page.height - margin_y
                for word in page.extract_words()
            ):
                edge_text = True
                break

    return {
        "page_count": page_count,
        "font_names": font_names,
        "tiny_char_fraction": tiny_chars / total_chars if total_chars else 0.0,
        "image_count": len(image_xrefs),
        "edge_text": edge_text,
        "is_encrypted": is_encrypted,
    }


def extract_docx(data):
    """Return (text, table_count) from a DOCX in one pass."""
    document = docx.Document(io.BytesIO(data))
    text = "\n".join(p.text for p in document.paragraphs)
    return text, len(document.tables)
