
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import re


def parse_markdown_line(line: str):
    original = line

    parts = []
    last_end = 0

    pattern = r'(\*\*.*?\*\*|\[.*?\]\(.*?\))'

    for match in re.finditer(pattern, line):
        if match.start() > last_end:
            parts.append(('normal', line[last_end:match.start()]))

        matched_text = match.group(0)

        if matched_text.startswith('**') and matched_text.endswith('**'):
            parts.append(('bold', matched_text[2:-2]))
        elif matched_text.startswith('['):
            link_match = re.match(r'\[(.*?)\]\((.*?)\)', matched_text)
            if link_match:
                parts.append(('link', link_match.group(1), link_match.group(2)))

        last_end = match.end()

    if last_end < len(line):
        parts.append(('normal', line[last_end:]))

    if not parts:
        parts = [('normal', original)]

    return parts


def markdown_to_pdf(markdown_text: str) -> bytes:
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.5 * inch,
        leftMargin=0.5 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch
    )

    story = []

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor='#000000',
        spaceAfter=8,
        alignment=TA_LEFT
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=13,
        textColor='#000000',
        spaceAfter=6,
        spaceBefore=12,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        textColor='#000000',
        spaceAfter=4,
        alignment=TA_LEFT,
        fontName='Helvetica'
    )

    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=11,
        textColor='#000000',
        spaceAfter=2,
        leftIndent=20,
        fontName='Helvetica'
    )

    lines = markdown_text.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            story.append(Spacer(1, 0.1 * inch))
            continue

        if line.startswith('# '):
            text = line[2:].strip()
            story.append(Paragraph(text, title_style))

        elif line.startswith('## '):
            text = line[3:].strip()
            story.append(Paragraph(text, heading_style))

        elif line.startswith('- '):
            text = line[2:].strip()

            parts = parse_markdown_line(text)
            formatted_text = ""

            for part in parts:
                if part[0] == 'bold':
                    formatted_text += f"<b>{part[1]}</b>"
                elif part[0] == 'link':
                    formatted_text += f'<font color="blue"><u>{part[1]}</u></font>'
                else:
                    formatted_text += part[1]

            story.append(Paragraph(formatted_text, bullet_style))

        else:
            parts = parse_markdown_line(line)
            formatted_text = ""

            for part in parts:
                if part[0] == 'bold':
                    formatted_text += f"<b>{part[1]}</b>"
                elif part[0] == 'link':
                    formatted_text += f'<font color="blue"><u>{part[1]}</u></font>'
                else:
                    formatted_text += part[1]

            story.append(Paragraph(formatted_text, normal_style))

    doc.build(story)

    buffer.seek(0)
    return buffer.read()


def markdown_to_docx(markdown_text: str) -> BytesIO:
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    lines = markdown_text.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if line.startswith('# '):
            text = line[2:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(24)
            run.font.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif line.startswith('## '):
            doc.add_paragraph()
            text = line[3:].strip()
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(13)
            run.font.bold = True
            para.paragraph_format.space_after = Pt(6)

        elif line.startswith('- '):
            text = line[2:].strip()

            para = doc.add_paragraph(style='List Bullet')

            parts = parse_markdown_line(text)

            for part in parts:
                if part[0] == 'bold':
                    run = para.add_run(part[1])
                    run.font.bold = True
                elif part[0] == 'link':
                    run = para.add_run(part[1])
                    run.font.color.rgb = RGBColor(0, 102, 204)
                else:
                    para.add_run(part[1])

            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(2)

        else:
            para = doc.add_paragraph()

            parts = parse_markdown_line(line)

            for part in parts:
                if part[0] == 'bold':
                    run = para.add_run(part[1])
                    run.font.bold = True
                elif part[0] == 'link':
                    run = para.add_run(part[1])
                    run.font.color.rgb = RGBColor(0, 102, 204)
                else:
                    para.add_run(part[1])

            para.paragraph_format.space_after = Pt(4)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def get_filename_base(markdown_text: str) -> str:
    lines = markdown_text.split('\n')
    for line in lines:
        if line.startswith('# '):
            name = line[2:].strip()

            name = re.sub(r'[^\w\s-]', '', name)
            name = re.sub(r'[-\s]+', '_', name)
            return name
    return "resume"